"""
Data Organizer Module - Step 1 of the preprocessing pipeline.

This module handles initial data organization and standardization:
- Merge Mz and RT columns into Mz/RT format (mz/RT)
- Simplify column headers (extract sample names from paths)
- Add Sample_Type row with auto-detection
- Parse method file (Word) for sample type mapping and injection sequence
- Create SampleInfo worksheet with injection order
- Reorder columns based on Injection_Order

Input format expected:
    Mz | RT | Intensity of path1 | Intensity of path2 | ...

Output format (RawIntensity):
    Mz/RT          | Sample1 | Sample2 | ...  (ordered by Injection_Order)
    Sample_Type    | exposure| control | ...
    252.1098/18.45 | 12345   | 67890   | ...

Output format (SampleInfo):
    Sample_Name | Sample_Type | Injection_Order | Injection_Volume
"""

import logging
import re
import zipfile
import xml.etree.ElementTree as ET
from pathlib import Path
from typing import Optional, Dict, Any, List, Tuple, Union
from dataclasses import dataclass
import pandas as pd
import numpy as np

logger = logging.getLogger(__name__)

from ms_core.preprocessing.base import BaseProcessor, ProcessingResult
from ms_core.preprocessing.settings import DataOrganizerConfig
from ms_core.utils.validators import detect_fixed_columns


@dataclass
class InjectionInfo:
    """Information about a sample injection from method file."""
    injection_order: int
    file_name: str
    sample_name: str
    injection_volume: float
    instrument_method: str = ""


class DataOrganizer(BaseProcessor):
    """
    Organizes and standardizes raw mass spectrometry data.

    This processor prepares the data for subsequent processing steps by:
    1. Merging Mz and RT columns into Mz/RT format
    2. Simplifying column headers to extract sample names
    3. Inserting Sample_Type row with auto-detection
    4. Optionally parsing method files for accurate sample mapping
    """

    # Patterns for sample type detection
    # Output mapping: tumor->Exposure, benign->Control, normal->Normal, qc->QC
    SAMPLE_TYPE_PATTERNS = {
        "QC": [r"qc", r"pool", r"quality"],
        "Exposure": [r"tumor", r"cancer", r"tumour"],  # Tumor -> Exposure
        "Normal": [r"normal", r"healthy"],
        "Control": [r"benign", r"benignfat"],  # Benign -> Control
        "blank": [r"blank", r"blk"],
        "standard": [r"std", r"standard", r"sdolek"],
    }
    SAMPLE_TOKEN_REGEX = re.compile(
        r"(?:EC\d{2,4}(?:_\d+)?|U\d{5}ZBEE|ZBEE\d{6}|pooled[_\s-]*QC[_\s-]*\d+|QC[_\s-]*sample[_\s-]*\d+|QC[_\s-]*\d+|blank)",
        re.IGNORECASE,
    )

    def __init__(self, config: Optional[DataOrganizerConfig] = None):
        """
        Initialize the Data Organizer.

        Args:
            config: Configuration options for data organization
        """
        super().__init__("Data Organizer")
        self.config = config or DataOrganizerConfig()

    def _is_non_sample_column(self, column_name: str) -> bool:
        """Return True when a column is metadata and should not be treated as a sample."""
        name = str(column_name).strip().lower()
        compact = re.sub(r"[^a-z0-9]+", "", name)
        if compact in {"rowid", "featureid", "feature", "id", "mzmineid", "mzminertmin", "z"}:
            return True
        if name.startswith("unnamed:"):
            return True
        if "mzmine rt" in name:
            return True
        return False

    def _normalize_sample_type_value(self, value: Any) -> Optional[str]:
        """Normalize sample type labels to toolkit's canonical values."""
        if value is None or pd.isna(value):
            return None
        text = str(value).strip()
        if not text:
            return None

        key = re.sub(r"[^a-z0-9]+", "", text.lower())
        mapping = {
            "qc": "QC",
            "qualitycontrol": "QC",
            "pooledqc": "QC",
            "exposure": "Exposure",
            "tumor": "Exposure",
            "tumour": "Exposure",
            "cancer": "Exposure",
            "case": "Exposure",
            "normal": "Normal",
            "control": "Control",
            "benign": "Control",
            "benignfat": "Control",
            "sample": "sample",
            "blank": "blank",
            "std": "standard",
            "standard": "standard",
            "sdolek": "standard",
            "na": "na",
        }
        return mapping.get(key, text)

    def _extract_sample_type_row_from_input(
        self,
        df: pd.DataFrame,
    ) -> Tuple[pd.DataFrame, Dict[str, str], Dict[str, Any]]:
        """
        Extract user-provided Sample type row from raw input when present.

        Expected marker is in the first row / first column (e.g., "Sample Type").
        """
        stats: Dict[str, Any] = {
            "sample_types_from_input": False,
            "input_sample_type_count": 0,
        }
        if df.empty:
            return df, {}, stats

        marker = str(df.iloc[0, 0]).strip().lower()
        marker_compact = re.sub(r"[^a-z0-9]+", "", marker)
        if marker_compact != "sampletype":
            return df, {}, stats

        provided_types: Dict[str, str] = {}
        for col in df.columns[2:]:
            col_str = str(col)
            if self._is_non_sample_column(col_str):
                continue
            normalized = self._normalize_sample_type_value(df.iloc[0][col])
            if normalized is None:
                continue
            provided_types[col_str] = normalized

        cleaned_df = df.iloc[1:].reset_index(drop=True)
        stats["sample_types_from_input"] = True
        stats["input_sample_type_count"] = len(provided_types)
        return cleaned_df, provided_types, stats

    def _extract_primary_sample_token(self, text: str) -> Optional[str]:
        """Extract a canonical sample token from free text."""
        if text is None:
            return None
        cleaned = str(text).strip()
        if not cleaned:
            return None
        match = self.SAMPLE_TOKEN_REGEX.search(cleaned)
        if not match:
            # BC-style tissue sample naming from breast-cancer method files.
            # Keep original wording so downstream BC/QC matching can use full context.
            compact = re.sub(r"\s+", "", cleaned.lower())
            if re.search(r"bc\d+_dna(?:\+rna|andrna)", compact):
                return cleaned
            if re.search(
                r"\bbc\d+_(?:dna\s*(?:\+\s*|and\s*)rna|dnaandrna|dna\+rna|dna|rna)\b",
                cleaned,
                re.IGNORECASE,
            ):
                return cleaned
            # Column-style names such as DNA_program1_TumorBC2257_DNA.
            if re.search(r"\b(?:dna_)?program\d+_[a-z0-9_]*bc\d+_(?:dnaandrna|dna|rna)\b", cleaned, re.IGNORECASE):
                return cleaned
            return None
        token = re.sub(r"\s+", "", match.group(0))
        token = token.replace("-", "_")
        return token

    def _normalize_sample_key(self, sample_name: str) -> str:
        """Normalize sample names from files/headers for robust matching."""
        if sample_name is None:
            return ""
        token = str(sample_name).strip()
        if not token:
            return ""
        token = self._extract_sample_name(token)
        lower = token.lower()
        lower = re.sub(r"[^a-z0-9]+", "_", lower).strip("_")
        lower = re.sub(r"_+", "_", lower)

        # Normalize QC naming variants.
        lower = re.sub(r"qc[_\s-]*sample[_\s-]*(\d+)", r"qc_sample_\1", lower)
        lower = re.sub(r"^qc[_\s-]*(\d+)$", r"qc_sample_\1", lower)
        lower = re.sub(r"pooled[_\s-]*qc[_\s-]*(\d+)", r"pooled_qc_\1", lower)

        # Normalize technical prefixes in column exports.
        lower = re.sub(r"^(?:dna|rna)_program\d+_", "", lower)
        lower = re.sub(r"^program\d+_", "", lower)
        lower = re.sub(r"dna_(?:and_)?rna", "dnaandrna", lower)
        lower = re.sub(r"dna_rna", "dnaandrna", lower)

        # Normalize ZBEE000070 -> U00070ZBEE.
        zbee_match = re.fullmatch(r"zbee(\d{6})", lower)
        if zbee_match:
            lower = f"u{int(zbee_match.group(1)):05d}zbee"

        # Normalize EC301 -> EC0301 and EC013_2 -> EC013.
        ec_suffix_match = re.fullmatch(r"(ec\d{2,4})_\d+", lower)
        if ec_suffix_match:
            lower = ec_suffix_match.group(1)
        ec_match = re.fullmatch(r"ec(\d{2,4})", lower)
        if ec_match:
            digits = ec_match.group(1)
            if len(digits) == 3 and int(digits) >= 300:
                lower = f"ec0{digits}"
            elif len(digits) == 2:
                lower = f"ec0{digits}"
            else:
                lower = f"ec{digits}"

        return lower

    def _is_likely_sample_name(self, text: str) -> bool:
        """Identify whether a method-file entry looks like a real sample."""
        if text is None:
            return False
        name_lower = str(text).lower()
        if "blank" in name_lower or "sdolek" in name_lower or "std" in name_lower:
            return False
        if self._extract_primary_sample_token(str(text)):
            return True
        return (
            re.search(r"bc\d+", name_lower) is not None
            or "qc" in name_lower
            or "tissue" in name_lower
            or "pooled" in name_lower
        )

    def validate_input(self, df: pd.DataFrame) -> tuple:
        """
        Validate input data for organization.

        Args:
            df: Input DataFrame

        Returns:
            Tuple of (is_valid, error_message)
        """
        if df is None or df.empty:
            return False, "Input data is empty"

        if len(df.columns) < 3:
            return False, "Input data must have at least 3 columns (Mz, RT, and at least one sample)"

        # Check if first two columns look like Mz and RT
        first_col = str(df.columns[0]).lower()
        second_col = str(df.columns[1]).lower()

        if "mz" not in first_col and "m/z" not in first_col and "mass" not in first_col:
            return False, f"First column '{df.columns[0]}' doesn't appear to be m/z values"

        if "rt" not in second_col and "time" not in second_col and "retention" not in second_col:
            return False, f"Second column '{df.columns[1]}' doesn't appear to be RT values"

        return True, ""

    def process(
        self,
        df: pd.DataFrame,
        method_file: Optional[Union[str, Path]] = None,
        mz_decimals: int = 4,
        rt_decimals: int = 2,
        sample_type_mapping: Optional[Dict[str, str]] = None,
        mode: str = "normalization",
        **kwargs,
    ) -> ProcessingResult:
        """
        Process raw data for organization.

        Args:
            df: Input DataFrame with Mz, RT, and intensity columns
            method_file: Optional path to method file (Word) for sample mapping
            mz_decimals: Decimal places for m/z values (default: 4)
            rt_decimals: Decimal places for RT values (default: 2)
            sample_type_mapping: Custom sample name to type mapping
            **kwargs: Additional parameters

        Returns:
            ProcessingResult with organized data (RawIntensity and SampleInfo)
        """
        self.reset()

        mode_name = str(mode or "normalization").strip().lower()
        if mode_name not in {"normalization", "statistics"}:
            return ProcessingResult(
                success=False,
                errors=[f"Unsupported mode: {mode}"],
                message=f"Unsupported mode: {mode}",
            )

        if mode_name == "statistics":
            return self._process_statistics_mode(
                df=df,
                method_file=method_file,
                mz_decimals=mz_decimals,
                rt_decimals=rt_decimals,
                sample_type_mapping=sample_type_mapping,
            )

        # Validate input
        is_valid, error_msg = self.validate_input(df)
        if not is_valid:
            return ProcessingResult(
                success=False,
                errors=[error_msg],
                message=f"Validation failed: {error_msg}",
            )

        self.update_progress(5, "Starting data organization...")

        try:
            # Create a copy
            result_df = df.copy()
            stats = {
                "original_rows": len(df),
                "original_cols": len(df.columns),
            }
            result_df, input_sample_types, input_type_stats = self._extract_sample_type_row_from_input(result_df)
            stats.update(input_type_stats)

            # Step 1: Parse method file if provided
            self.update_progress(10, "Parsing method file...")
            sample_mapping = {}
            injection_info_list: List[InjectionInfo] = []

            if method_file:
                sample_mapping = self._parse_method_file(method_file)
                injection_info_list = self._parse_injection_sequence(method_file)
                stats["method_file_samples"] = len(sample_mapping)
                stats["injection_sequence_count"] = len(injection_info_list)

            if sample_type_mapping:
                sample_mapping.update(sample_type_mapping)

            if self._cancelled:
                return ProcessingResult(success=False, message="Processing cancelled")

            # Step 2: Merge Mz and RT into Mz/RT
            self.update_progress(30, "Merging Mz/RT columns...")
            result_df, merge_stats = self._merge_mz_rt(
                result_df, mz_decimals, rt_decimals, include_tolerance_col=False
            )
            stats.update(merge_stats)

            if self._cancelled:
                return ProcessingResult(success=False, message="Processing cancelled")

            # Step 3: Simplify column headers
            self.update_progress(50, "Simplifying column headers...")
            result_df, header_mapping = self._simplify_headers(result_df)
            stats["columns_simplified"] = len(header_mapping)
            input_sample_types_simplified: Dict[str, str] = {}
            if input_sample_types:
                for raw_col, sample_type in input_sample_types.items():
                    simplified = header_mapping.get(raw_col, self._extract_sample_name(str(raw_col)))
                    normalized = self._normalize_sample_type_value(sample_type)
                    if simplified and normalized and simplified not in input_sample_types_simplified:
                        input_sample_types_simplified[simplified] = normalized

            if self._cancelled:
                return ProcessingResult(success=False, message="Processing cancelled")

            # Step 4: Insert Sample_Type row
            self.update_progress(60, "Detecting sample types...")
            result_df, type_stats = self._insert_sample_type_row(
                result_df,
                header_mapping,
                sample_mapping,
                sample_type_overrides=input_sample_types_simplified,
            )
            stats.update(type_stats)

            if self._cancelled:
                return ProcessingResult(success=False, message="Processing cancelled")

            # Step 5: Build SampleInfo DataFrame
            self.update_progress(70, "Building SampleInfo...")
            sample_info_df = self._build_sample_info(
                result_df, injection_info_list
            )
            stats["sample_info_rows"] = len(sample_info_df)

            if self._cancelled:
                return ProcessingResult(success=False, message="Processing cancelled")

            # Step 6: Reorder columns based on Injection_Order
            self.update_progress(80, "Reordering columns by injection order...")
            result_df = self._reorder_columns_by_injection(result_df, sample_info_df)

            # Clean up SampleInfo: remove internal columns
            if "_col_name" in sample_info_df.columns:
                sample_info_df = sample_info_df.drop(columns=["_col_name"])

            if self._cancelled:
                return ProcessingResult(success=False, message="Processing cancelled")

            # Step 7: Final cleanup
            self.update_progress(90, "Finalizing...")
            result_df = self._finalize_structure(result_df)

            self.update_progress(100, "Data organization complete")

            stats["final_rows"] = len(result_df)
            stats["final_cols"] = len(result_df.columns)

            return ProcessingResult(
                success=True,
                data=result_df,
                message=f"Data organization completed. {stats['original_rows']} features processed.",
                statistics=stats,
                metadata={
                    "mode": mode_name,
                    "header_mapping": header_mapping,
                    "sample_mapping": sample_mapping,
                    "sample_info": sample_info_df,
                },
            )

        except Exception as e:
            return ProcessingResult(
                success=False,
                errors=[str(e)],
                message=f"Error during data organization: {str(e)}",
            )

    def _process_statistics_mode(
        self,
        df: pd.DataFrame,
        method_file: Optional[Union[str, Path]] = None,
        mz_decimals: int = 4,
        rt_decimals: int = 2,
        sample_type_mapping: Optional[Dict[str, str]] = None,
    ) -> ProcessingResult:
        """
        Statistics mode:
        - Keep separate Mz and RT output columns (no merged Mz/RT in final output)
        - Otherwise follow normalization workflow
        """
        is_valid, error_msg = self.validate_input(df)
        if not is_valid:
            return ProcessingResult(
                success=False,
                errors=[error_msg],
                message=f"Validation failed: {error_msg}",
            )

        self.update_progress(5, "Starting statistics mode...")

        try:
            result_df = df.copy()
            stats: Dict[str, Any] = {
                "mode": "statistics",
                "original_rows": len(df),
                "original_cols": len(df.columns),
            }
            result_df, input_sample_types, input_type_stats = self._extract_sample_type_row_from_input(result_df)
            stats.update(input_type_stats)
            original_mz_col = str(df.columns[0])
            original_rt_col = str(df.columns[1])
            original_mz_values = result_df.iloc[:, 0].tolist()
            original_rt_values = result_df.iloc[:, 1].tolist()

            # Step 1: Parse method file if provided
            self.update_progress(10, "Parsing method file...")
            sample_mapping: Dict[str, str] = {}
            injection_info_list: List[InjectionInfo] = []
            if method_file:
                sample_mapping = self._parse_method_file(method_file)
                injection_info_list = self._parse_injection_sequence(method_file)
            if sample_type_mapping:
                sample_mapping.update(sample_type_mapping)
            stats["method_file_samples"] = len(sample_mapping)
            stats["injection_sequence_count"] = len(injection_info_list)

            if self._cancelled:
                return ProcessingResult(success=False, message="Processing cancelled")

            # Step 2: Keep normalization internals, but restore separate Mz/RT later
            self.update_progress(30, "Preparing internal Mz/RT for downstream steps...")
            result_df, merge_stats = self._merge_mz_rt(
                result_df,
                mz_decimals=mz_decimals,
                rt_decimals=rt_decimals,
                include_tolerance_col=False,
            )
            stats["merge_skipped_in_output"] = True
            stats["invalid_values"] = merge_stats.get("invalid_values", 0)
            stats["mz_rt_merged"] = 0

            if self._cancelled:
                return ProcessingResult(success=False, message="Processing cancelled")

            # Step 3: Simplify column headers
            self.update_progress(50, "Simplifying column headers...")
            result_df, header_mapping = self._simplify_headers(result_df)
            stats["columns_simplified"] = len(header_mapping)
            input_sample_types_simplified: Dict[str, str] = {}
            if input_sample_types:
                for raw_col, sample_type in input_sample_types.items():
                    simplified = header_mapping.get(raw_col, self._extract_sample_name(str(raw_col)))
                    normalized = self._normalize_sample_type_value(sample_type)
                    if simplified and normalized and simplified not in input_sample_types_simplified:
                        input_sample_types_simplified[simplified] = normalized

            if self._cancelled:
                return ProcessingResult(success=False, message="Processing cancelled")

            # Step 4: Insert Sample_Type row
            self.update_progress(60, "Detecting sample types...")
            result_df, type_stats = self._insert_sample_type_row(
                result_df,
                header_mapping,
                sample_mapping,
                sample_type_overrides=input_sample_types_simplified,
            )
            stats.update(type_stats)

            if self._cancelled:
                return ProcessingResult(success=False, message="Processing cancelled")

            # Step 5: Build SampleInfo DataFrame
            self.update_progress(70, "Building SampleInfo...")
            sample_info_df = self._build_sample_info(
                result_df, injection_info_list
            )
            stats["sample_info_rows"] = len(sample_info_df)

            if self._cancelled:
                return ProcessingResult(success=False, message="Processing cancelled")

            # Step 6: Reorder columns based on Injection_Order
            self.update_progress(80, "Reordering columns by injection order...")
            result_df = self._reorder_columns_by_injection(result_df, sample_info_df)

            # Clean up SampleInfo: remove internal columns
            if "_col_name" in sample_info_df.columns:
                sample_info_df = sample_info_df.drop(columns=["_col_name"])

            if self._cancelled:
                return ProcessingResult(success=False, message="Processing cancelled")

            # Step 7: Final cleanup
            self.update_progress(90, "Finalizing...")
            result_df = self._finalize_structure(result_df)

            # Step 8: Restore separate Mz/RT columns in final output
            self.update_progress(95, "Restoring separate Mz and RT columns...")
            mz_column_values = ["Sample_Type"] + original_mz_values
            rt_column_values = ["na"] + original_rt_values
            result_df.insert(0, original_mz_col, mz_column_values)
            result_df.insert(1, original_rt_col, rt_column_values)

            # Remove internal merged Mz/RT column while preserving all other columns.
            mzrt_positions = [idx for idx, col in enumerate(result_df.columns) if col == "Mz/RT"]
            if mzrt_positions:
                drop_idx = mzrt_positions[0]
                keep_idx = [idx for idx in range(len(result_df.columns)) if idx != drop_idx]
                result_df = result_df.iloc[:, keep_idx]

            self.update_progress(100, "Statistics mode complete")
            stats["final_rows"] = len(result_df)
            stats["final_cols"] = len(result_df.columns)

            return ProcessingResult(
                success=True,
                data=result_df,
                message="Statistics mode completed.",
                statistics=stats,
                metadata={
                    "mode": "statistics",
                    "header_mapping": header_mapping,
                    "sample_mapping": sample_mapping,
                    "sample_info": sample_info_df,
                },
            )
        except Exception as e:
            return ProcessingResult(
                success=False,
                errors=[str(e)],
                message=f"Error during statistics mode: {str(e)}",
            )

    def _validate_statistics_input(self, df: pd.DataFrame) -> Tuple[bool, str]:
        """Validate input data for statistics mode."""
        if df is None or df.empty:
            return False, "Input data is empty"
        if len(df.columns) < 3:
            return False, "Input data must have at least 3 columns"

        _, num_fixed = self._detect_fixed_columns_for_statistics(df)
        if num_fixed >= 1:
            return True, ""

        first_col = str(df.columns[0]).lower()
        second_col = str(df.columns[1]).lower()
        if (("mz" in first_col or "m/z" in first_col or "mass" in first_col)
                and ("rt" in second_col or "time" in second_col or "retention" in second_col)):
            return True, ""

        return False, (
            "Statistics mode expects either normalized fixed columns "
            "(Mz/RT) or raw Mz/RT leading columns"
        )

    def _detect_fixed_columns_for_statistics(self, df: pd.DataFrame) -> Tuple[List[str], int]:
        """Detect fixed columns for statistics-mode sorting."""
        fixed_cols, num_fixed = detect_fixed_columns(df)
        if num_fixed == 0:
            return fixed_cols, num_fixed

        # Accept legacy tolerance label as fixed column in statistics mode.
        if num_fixed < len(df.columns):
            next_col = str(df.columns[num_fixed]).lower()
            if "tolerance" in next_col:
                fixed_cols = fixed_cols + [df.columns[num_fixed]]
                num_fixed += 1

        return fixed_cols, num_fixed

    def _reorder_columns_statistics_mode(
        self,
        df: pd.DataFrame,
        injection_info_list: List[InjectionInfo],
    ) -> Tuple[pd.DataFrame, Dict[str, Any]]:
        """Reorder sample columns for statistics mode without creating SampleInfo."""
        stats: Dict[str, Any] = {
            "columns_reordered": 0,
            "columns_unmatched": 0,
        }
        if df.empty:
            return df, stats

        fixed_cols, num_fixed = self._detect_fixed_columns_for_statistics(df)
        fixed_positions = list(range(num_fixed))
        sample_positions = [
            idx for idx in range(num_fixed, len(df.columns))
            if not self._is_non_sample_column(str(df.columns[idx]))
        ]
        metadata_positions = [
            idx for idx in range(num_fixed, len(df.columns))
            if self._is_non_sample_column(str(df.columns[idx]))
        ]

        if not sample_positions:
            return df, stats

        filtered_injection_list: List[InjectionInfo] = []
        for info in injection_info_list:
            if self._is_likely_sample_name(info.file_name):
                filtered_injection_list.append(info)

        if not filtered_injection_list:
            stats["columns_unmatched"] = len(sample_positions)
            return df, stats

        sorted_injection_list = sorted(filtered_injection_list, key=lambda x: x.injection_order)

        available_positions = list(sample_positions)
        ordered_positions: List[int] = []

        for info in sorted_injection_list:
            match_pos = self._find_matching_sample_column_position(
                df,
                available_positions,
                info.file_name,
            )
            if match_pos is None:
                continue
            ordered_positions.append(match_pos)
            available_positions.remove(match_pos)

        ordered_positions.extend(available_positions)
        stats["columns_reordered"] = len(ordered_positions) - len(available_positions)
        stats["columns_unmatched"] = len(available_positions)

        new_positions = fixed_positions + ordered_positions + metadata_positions
        reordered_df = df.iloc[:, new_positions]
        reordered_df.columns = (
            fixed_cols
            + [df.columns[i] for i in ordered_positions]
            + [df.columns[i] for i in metadata_positions]
        )
        return reordered_df, stats

    def _find_matching_sample_column_position(
        self,
        df: pd.DataFrame,
        candidate_positions: List[int],
        file_name: str,
    ) -> Optional[int]:
        """Find a matching sample column index for a method-file sample name."""
        file_lower = file_name.lower().replace("\n", " ")
        file_simplified = self._simplify_word_sample_name(file_name).lower()
        file_keys = {
            self._normalize_sample_key(file_name),
            self._normalize_sample_key(file_simplified),
            self._normalize_sample_key(self._extract_primary_sample_token(file_name) or ""),
        }
        file_keys = {k for k in file_keys if k}

        def detect_variant(text: str) -> str:
            text = text.lower().replace("\n", " ").replace("*", " ")
            if "dna" in text and "rna" in text and "+" in text:
                return "dna+rna"
            if "dnaandrna" in text.replace(" ", ""):
                return "dna+rna"
            if "_rna" in text or " rna" in text or text.endswith("rna"):
                return "rna"
            return "dna"

        def sanitize(text: str) -> str:
            return re.sub(r"[^a-z0-9]+", "", text.lower())

        for pos in candidate_positions:
            col_raw = str(df.columns[pos])
            col_lower = col_raw.lower()
            col_key = self._extract_sample_name(col_raw).lower()
            col_keys = {
                self._normalize_sample_key(col_raw),
                self._normalize_sample_key(col_key),
                self._normalize_sample_key(self._extract_primary_sample_token(col_raw) or ""),
            }
            col_keys = {k for k in col_keys if k}

            if file_keys and col_keys and file_keys.intersection(col_keys):
                return pos

            bc_match_col = re.search(r"(tumor|normal|benign|benignfat)?(bc\d+)", col_key)
            bc_match_file = re.search(r"(tumor|normal|benign)\s*(tissue)?\s*(fat\s*)?(bc\d+)", file_lower)
            if bc_match_col and bc_match_file:
                col_prefix = bc_match_col.group(1) or ""
                if "benign" in col_prefix:
                    col_prefix = "benign"
                col_id = bc_match_col.group(2)
                col_variant = detect_variant(col_key)

                file_prefix = bc_match_file.group(1) or ""
                file_id = bc_match_file.group(4)
                file_variant = detect_variant(file_lower)
                if col_id == file_id and col_prefix == file_prefix and col_variant == file_variant:
                    return pos

            qc_match_col = re.search(r"(pooled_?)?qc_?(\d+)", col_key)
            qc_match_file = re.search(r"(pooled_?)?qc_?(\d+)", file_lower)
            if qc_match_col and qc_match_file and qc_match_col.group(2) == qc_match_file.group(2):
                return pos

            if file_simplified and (file_simplified in col_key or col_key in file_simplified):
                return pos

            col_token = sanitize(col_key)
            file_token = sanitize(file_simplified)
            if col_token and file_token and (col_token == file_token or col_token in file_token or file_token in col_token):
                return pos

            if file_simplified and (file_simplified in col_lower or col_lower in file_simplified):
                return pos

        return None

    def _merge_mz_rt(
        self,
        df: pd.DataFrame,
        mz_decimals: int = 4,
        rt_decimals: int = 2,
        include_tolerance_col: bool = True,
    ) -> Tuple[pd.DataFrame, Dict[str, Any]]:
        """
        Merge Mz and RT columns into a single Mz/RT column.

        Format: "mz/RT" (e.g., "252.1098/18.45")
        """
        stats = {"mz_rt_merged": 0, "invalid_values": 0}

        # Get Mz and RT columns (first two columns)
        mz_col = df.columns[0]
        rt_col = df.columns[1]

        # Vectorized merge for performance
        mz_series = pd.to_numeric(df.iloc[:, 0], errors="coerce")
        rt_series = pd.to_numeric(df.iloc[:, 1], errors="coerce")
        valid_mask = mz_series.notna() & rt_series.notna()

        mz_np = mz_series.to_numpy()
        rt_np = rt_series.to_numpy()
        mz_str = np.char.mod(f"%.{mz_decimals}f", mz_np)
        rt_str = np.char.mod(f"%.{rt_decimals}f", rt_np)
        merged = np.char.add(np.char.add(mz_str, "/"), rt_str)

        # Fallback to original strings for invalid rows
        orig_mz = df.iloc[:, 0].astype(str).to_numpy()
        orig_rt = df.iloc[:, 1].astype(str).to_numpy()
        fallback = np.char.add(np.char.add(orig_mz, "/"), orig_rt)
        mz_rt_values = np.where(valid_mask.to_numpy(), merged, fallback).tolist()

        stats["mz_rt_merged"] = int(valid_mask.sum())
        stats["invalid_values"] = int(len(df) - valid_mask.sum())

        # Build new frame in one concat to avoid highly-fragmented DataFrame writes.
        front_cols = {"Mz/RT": mz_rt_values}
        if include_tolerance_col:
            front_cols["m/z Tolerance( ppm)/RT Tolerance"] = "na"
        leading_df = pd.DataFrame(front_cols)
        trailing_df = df.iloc[:, 2:].reset_index(drop=True)
        new_df = pd.concat([leading_df, trailing_df], axis=1)

        return new_df, stats

    def _simplify_headers(
        self,
        df: pd.DataFrame,
    ) -> Tuple[pd.DataFrame, Dict[str, str]]:
        """
        Simplify column headers by extracting sample names from paths.

        Examples:
            "Intensity of C:\\...\\program2_program1_TumorBC2257_DNA.tsv" -> "TumorBC2257_DNA"
            "Intensity of C:\\...\\Breast_Cancer_Tissue_pooled_QC1.tsv" -> "pooled_QC1"
        """
        header_mapping = {}  # old_name -> new_name

        # Fixed column names that should not be simplified
        fixed_cols = ["Mz/RT", "FeatureID", "m/z Tolerance( ppm)/RT Tolerance"]

        new_columns = []
        for col in df.columns:
            col_str = str(col)

            # Skip fixed columns
            if col in fixed_cols:
                new_columns.append(col)
                continue

            # Extract sample name from path
            new_name = self._extract_sample_name(col_str)
            header_mapping[col_str] = new_name
            new_columns.append(new_name)

        df.columns = new_columns
        return df, header_mapping

    def _extract_sample_name(self, header: str) -> str:
        """
        Extract sample name from a column header (typically a file path).

        Args:
            header: Original column header

        Returns:
            Simplified sample name
        """
        # Remove "Intensity of " prefix if present
        if header.lower().startswith("intensity of "):
            header = header[13:]

        # Get the filename from the path
        try:
            # Handle both Windows and Unix paths
            path = Path(header)
            filename = path.stem  # Get filename without extension
        except Exception as exc:
            logger.debug("Path parse fallback for header '%s': %s", header, exc)
            filename = header

        # Try to extract the meaningful part
        # Pattern: program2_program1_SAMPLENAME.tsv -> SAMPLENAME
        # Pattern: program2_1\\program2_program1_SAMPLENAME -> SAMPLENAME
        patterns_to_remove = [
            r"^program\d+_(?:dna|rna)_program\d+_",  # program2_DNA_program1_
            r"^(?:dna|rna)_program\d+_",  # DNA_program1_, RNA_program1_
            r"^program\d+_program\d+_",  # program2_program1_
            r"^program\d+_\d+_",  # program2_1_
            r"^program\d+_",  # program2_
        ]

        result = filename
        for pattern in patterns_to_remove:
            result = re.sub(pattern, "", result, flags=re.IGNORECASE)

        # Normalize QC naming: QC4 -> QC_4 (standardize with underscore)
        result = re.sub(r"(qc)[ _-]?(\d+)", r"\1_\2", result, flags=re.IGNORECASE)

        # Clean up any remaining artifacts
        result = result.strip("_")

        return result if result else filename

    def _insert_sample_type_row(
        self,
        df: pd.DataFrame,
        header_mapping: Dict[str, str],
        sample_mapping: Dict[str, str],
        sample_type_overrides: Optional[Dict[str, str]] = None,
    ) -> Tuple[pd.DataFrame, Dict[str, Any]]:
        """
        Insert Sample_Type row at the top of the data.

        Auto-detects sample types based on column names, unless user-provided
        sample types are available from the input file.
        """
        stats = {"types_detected": {}, "types_from_input": 0}

        # Determine number of fixed columns (Mz/RT only, or Mz/RT + Tolerance)
        fixed_cols, num_fixed = detect_fixed_columns(df)

        # Detect sample types
        sample_types = ["Sample_Type"] + ["na"] * (num_fixed - 1)  # Fixed columns
        override_exact: Dict[str, str] = {}
        override_by_key: Dict[str, str] = {}
        if sample_type_overrides:
            for col_name, sample_type in sample_type_overrides.items():
                normalized = self._normalize_sample_type_value(sample_type)
                if normalized is None:
                    continue
                exact_key = str(col_name)
                override_exact[exact_key] = normalized
                normalized_col_key = self._normalize_sample_key(exact_key)
                if normalized_col_key and normalized_col_key not in override_by_key:
                    override_by_key[normalized_col_key] = normalized

        for col in df.columns[num_fixed:]:
            sample_type = override_exact.get(str(col))
            if sample_type is None:
                col_key = self._normalize_sample_key(str(col))
                if col_key:
                    sample_type = override_by_key.get(col_key)
            if sample_type is None:
                sample_type = self._detect_sample_type(col, sample_mapping)
            else:
                stats["types_from_input"] += 1
            sample_types.append(sample_type)

            # Track statistics
            if sample_type not in stats["types_detected"]:
                stats["types_detected"][sample_type] = 0
            stats["types_detected"][sample_type] += 1

        # Create sample type row as DataFrame
        sample_type_row = pd.DataFrame([sample_types], columns=df.columns)

        # Concatenate: sample_type_row on top of data
        result_df = pd.concat([sample_type_row, df], ignore_index=True)

        return result_df, stats

    def _detect_sample_type(
        self,
        column_name: str,
        sample_mapping: Dict[str, str],
    ) -> str:
        """
        Detect sample type from column name.

        Priority:
        1. Direct pattern matching from column name (e.g., TumorBC2257 -> tumor)
        2. Pre-defined mapping from method file (fallback)
        3. Default to "sample"

        Args:
            column_name: Name of the column (sample name)
            sample_mapping: Pre-defined mapping from method file

        Returns:
            Detected sample type
        """
        col_lower = column_name.lower()

        # Priority 1: Direct pattern matching from column name
        # This ensures "TumorBC2257_DNA" is detected as tumor, not normal
        for sample_type, patterns in self.SAMPLE_TYPE_PATTERNS.items():
            for pattern in patterns:
                if re.search(pattern, col_lower, re.IGNORECASE):
                    return sample_type

        # Priority 2: Check pre-defined mapping (only if direct detection fails)
        for pattern, sample_type in sample_mapping.items():
            if pattern.lower() in col_lower:
                return sample_type

        # Default to "sample" if no pattern matches
        return "sample"

    def _finalize_structure(self, df: pd.DataFrame) -> pd.DataFrame:
        """
        Finalize the DataFrame structure.

        Ensures proper data types and formatting.
        """
        # Determine number of fixed columns
        fixed_cols, num_fixed = detect_fixed_columns(df)

        # First row is now Sample_Type — data rows start from index 1.
        # Vectorized numeric conversion for all sample columns at once.
        if len(df) > 1 and num_fixed < len(df.columns):
            # Use positional indexing to support duplicate sample column names.
            # Label-based assignment can fail when column labels are not unique.
            sample_col_positions = list(range(num_fixed, len(df.columns)))
            data_block = df.iloc[1:, sample_col_positions]
            converted = data_block.apply(pd.to_numeric, errors="coerce")
            conv_values = converted.to_numpy()
            # Rebuild each column as a Python list so pandas 3.x
            # StringDtype columns are replaced with object dtype.
            for j, col_pos in enumerate(sample_col_positions):
                col_name = df.columns[col_pos]
                df[col_name] = [df.iat[0, col_pos]] + conv_values[:, j].tolist()

        return df

    def _extract_docx_tables_fallback(self, file_path: Union[str, Path]) -> List[List[List[str]]]:
        """
        Extract DOCX table cell texts without python-docx.

        This fallback parses word/document.xml directly so the pipeline can still
        build SampleInfo when python-docx is unavailable.
        """
        tables: List[List[List[str]]] = []
        path = Path(file_path)

        if path.suffix.lower() != ".docx" or not path.exists():
            return tables

        try:
            with zipfile.ZipFile(path, "r") as zf:
                with zf.open("word/document.xml") as fp:
                    tree = ET.parse(fp)
            root = tree.getroot()
            ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}

            for tbl in root.findall(".//w:tbl", ns):
                parsed_rows: List[List[str]] = []
                for tr in tbl.findall("./w:tr", ns):
                    parsed_cells: List[str] = []
                    for tc in tr.findall("./w:tc", ns):
                        texts = [t.text for t in tc.findall(".//w:t", ns) if t.text]
                        parsed_cells.append("".join(texts).strip())
                    if parsed_cells:
                        parsed_rows.append(parsed_cells)
                if parsed_rows:
                    tables.append(parsed_rows)
        except Exception as exc:
            logger.warning("Fallback DOCX parse failed for %s: %s", path, exc)

        return tables

    def _parse_injection_volume_from_cells(self, cells: List[str]) -> float:
        """Parse likely injection volume from a table row."""
        for cell in reversed(cells):
            cell_text = str(cell).strip().replace(",", "")
            if not cell_text:
                continue
            try:
                value = float(cell_text)
            except ValueError:
                continue
            if 0 < value <= 100:
                return value
        return 0.0

    def _extract_injection_rows_from_table(
        self,
        table_rows: List[List[str]],
        preserve_source_order: bool = False,
    ) -> List[InjectionInfo]:
        """
        Extract injection rows from one table using common layouts.

        Supported layouts:
        - [Order, Sample, ...]
        - [Order, Sample, ..., Order, Sample, ...] (dual-column Word table)
        """
        injections: List[InjectionInfo] = []
        candidate_pairs = [(0, 1), (3, 4)]

        for row in table_rows:
            cells = [str(cell).strip() for cell in row]
            if len(cells) < 2:
                continue

            for order_idx, sample_idx in candidate_pairs:
                if len(cells) <= sample_idx:
                    continue
                order_text = cells[order_idx].strip() if len(cells) > order_idx else ""
                if not re.fullmatch(r"\d{1,4}", order_text):
                    continue

                sample_cell = cells[sample_idx]
                sample_token = self._extract_primary_sample_token(sample_cell)
                if not sample_token:
                    continue
                sample_text = re.sub(r"\s+", " ", sample_cell).strip()
                if not sample_text:
                    sample_text = sample_token

                instrument_method = ""
                next_col = sample_idx + 1
                if len(cells) > next_col:
                    candidate = re.sub(r"\s+", " ", cells[next_col]).strip()
                    if candidate and not candidate.isdigit():
                        instrument_method = candidate

                injection_volume = self._parse_injection_volume_from_cells(cells)
                sample_name = self._simplify_word_sample_name(sample_text)

                injections.append(
                    InjectionInfo(
                        injection_order=int(order_text),
                        file_name=sample_text,
                        sample_name=sample_name,
                        injection_volume=injection_volume,
                        instrument_method=instrument_method,
                    )
                )

        if not injections:
            return injections

        if preserve_source_order:
            deduped_source: List[InjectionInfo] = []
            seen_keys = set()
            for info in injections:
                key = self._normalize_sample_key(info.file_name)
                if not key or key in seen_keys:
                    continue
                seen_keys.add(key)
                deduped_source.append(info)
            for idx, info in enumerate(deduped_source, start=1):
                info.injection_order = idx
            return deduped_source

        deduped: List[InjectionInfo] = []
        seen = set()
        for info in sorted(injections, key=lambda x: x.injection_order):
            dedupe_key = (info.injection_order, self._normalize_sample_key(info.file_name))
            if dedupe_key in seen:
                continue
            seen.add(dedupe_key)
            deduped.append(info)
        return deduped

    def _parse_injection_sequence(self, file_path: Union[str, Path]) -> List[InjectionInfo]:
        """
        Parse injection sequence table from Word document.

        Extracts samples in row order (ignoring ID column which may have duplicates
        due to pagination issues in Word tables).

        Args:
            file_path: Path to the Word document

        Returns:
            List of InjectionInfo objects in injection order
        """
        injection_list: List[InjectionInfo] = []
        file_path = Path(file_path)

        if not file_path.exists():
            return injection_list

        if file_path.suffix.lower() not in [".docx", ".doc"]:
            return injection_list

        tables: List[List[List[str]]] = []
        try:
            from docx import Document
            doc = Document(file_path)
            for table in doc.tables:
                table_rows: List[List[str]] = []
                for row in table.rows:
                    table_rows.append([cell.text.strip() for cell in row.cells])
                if table_rows:
                    tables.append(table_rows)
        except ImportError:
            logger.warning("python-docx not installed; using fallback DOCX parser for injection sequence")
            tables = self._extract_docx_tables_fallback(file_path)
        except Exception as exc:
            logger.warning("python-docx parse failed for %s (%s); using fallback parser", file_path, exc)
            tables = self._extract_docx_tables_fallback(file_path)

        # Find the injection sequence table
        # Look for table with columns similar to: ID | File Name | Instrument Method | ...
        target_table = None
        for table_rows in tables:
            if len(table_rows) > 10:  # Must have many rows
                header_cells = [str(cell).strip().lower() for cell in table_rows[0]]
                if any(
                    ("file" in h and "name" in h)
                    or ("filename" in h)
                    or ("檔" in h)
                    or ("樣本" in h)
                    for h in header_cells
                ):
                    target_table = table_rows
                    break

        # Preferred parser: extract (order, sample) pairs directly from candidate table.
        if target_table is not None:
            parsed_from_target = self._extract_injection_rows_from_table(target_table)
            if parsed_from_target:
                # Some BC method sheets reuse numeric IDs across sections; in that case
                # source-row order is a better approximation of true injection sequence.
                order_values = [info.injection_order for info in parsed_from_target]
                duplicate_orders = len(order_values) - len(set(order_values))
                bc_like_count = sum(
                    1 for info in parsed_from_target
                    if re.search(r"bc\d+", info.file_name, re.IGNORECASE)
                )
                if duplicate_orders > 0 and bc_like_count >= 5:
                    parsed_by_rows = self._extract_injection_rows_from_table(
                        target_table,
                        preserve_source_order=True,
                    )
                    if parsed_by_rows:
                        return parsed_by_rows
                return parsed_from_target

        # Parse the table using ROW ORDER (not ID column)
        # This handles Word tables where IDs may reset across pages
        if target_table is not None:
            row_order = 0
            for row_idx, row in enumerate(target_table):
                if row_idx == 0:  # Skip header row
                    continue

                cells = [str(cell).strip() for cell in row]
                if len(cells) < 2:
                    continue

                # Parse file name (sample name) - column 1
                file_name = cells[1] if len(cells) > 1 else ""
                # Normalize whitespace/newlines from Word tables
                file_name = re.sub(r"\s+", " ", file_name).strip()
                # Keep original underscores; only normalize whitespace
                if not file_name:
                    continue

                # Use row order as injection order (will be renumbered later)
                row_order += 1

                # Parse instrument method (column 2 or 3)
                instrument_method = ""
                for i in [2, 3]:
                    if len(cells) > i and "method" not in cells[i].lower():
                        if cells[i] and not cells[i].isdigit():
                            instrument_method = cells[i]
                            break

                injection_volume = self._parse_injection_volume_from_cells(cells)
                sample_name = self._simplify_word_sample_name(file_name)

                injection_list.append(
                    InjectionInfo(
                        injection_order=row_order,  # Use row order
                        file_name=file_name,
                        sample_name=sample_name,
                        injection_volume=injection_volume,
                        instrument_method=instrument_method,
                    )
                )
            if injection_list:
                return injection_list

        # Fallback parser: score all tables and choose the one with the best
        # (order, sample) extraction coverage.
        best_list: List[InjectionInfo] = []
        best_score: Tuple[int, int, int, int] = (-1, -1, -1, -1)
        for table_rows in tables:
            parsed = self._extract_injection_rows_from_table(table_rows)
            if len(parsed) < 5:
                continue

            unique_orders = len({info.injection_order for info in parsed})
            unique_samples = len({self._normalize_sample_key(info.file_name) for info in parsed})
            duplicate_orders = len(parsed) - unique_orders
            score = (unique_samples, unique_orders, -duplicate_orders, len(parsed))
            if score > best_score:
                best_score = score
                best_list = parsed

        return best_list

    def _simplify_word_sample_name(self, file_name: str) -> str:
        """
        Simplify sample name from Word document to match column headers.

        Examples:
            "Breast Cancer Tissue_ pooled_QC_1" -> "pooled_QC_1"
            "Tumor tissue BC2257_DNA" -> "TumorBC2257_DNA"
            "Normal tissue BC2257_DNA" -> "NormalBC2257_DNA"
            "Benign tissue Fat BC2250_DNA" -> "BenignBC2250_DNA"

        Args:
            file_name: Original sample name from Word document

        Returns:
            Simplified sample name that matches column header
        """
        name = file_name.strip()

        # Handle QC samples: "Breast Cancer Tissue_ pooled_QC_X" -> "pooled_QC_X"
        if "pooled" in name.lower() and "qc" in name.lower():
            match = re.search(r"pooled_?QC_?\d+", name, re.IGNORECASE)
            if match:
                return match.group().replace(" ", "")

        # Handle tissue samples: "Tumor tissue BC2257_DNA" -> "TumorBC2257_DNA"
        tissue_patterns = [
            (r"tumor\s*tissue\s*", "Tumor"),
            (r"normal\s*tissue\s*", "Normal"),
            (r"benign\s*tissue\s*(fat\s*)?", "Benign"),
        ]

        for pattern, prefix in tissue_patterns:
            match = re.search(pattern, name, re.IGNORECASE)
            if match:
                # Extract the BC ID part
                bc_match = re.search(r"BC\d+_\w+", name, re.IGNORECASE)
                if bc_match:
                    return f"{prefix}{bc_match.group()}"

        # Handle blank and standard samples
        if "blank" in name.lower():
            match = re.search(r"blank_?\d*", name, re.IGNORECASE)
            if match:
                return match.group()

        if "sdolek" in name.lower() or "std" in name.lower():
            # Return as-is but cleaned
            return re.sub(r"\s+", "_", name)

        return name

    def _build_sample_info(
        self,
        df: pd.DataFrame,
        injection_info_list: List[InjectionInfo],
    ) -> pd.DataFrame:
        """
        Build SampleInfo DataFrame from column names and injection info.

        Creates a mapping between column names and injection order.
        Re-numbers Injection_Order starting from 1, excluding blanks and standards.

        Args:
            df: DataFrame with simplified column headers
            injection_info_list: List of InjectionInfo from method file

        Returns:
            SampleInfo DataFrame with columns:
            - Sample_Name: Full sample name from Word document
            - Sample_Type: Detected sample type
            - Injection_Order: Re-numbered order (starting from 1)
            - Injection_Volume: Volume from method file
        """
        # Determine number of fixed columns
        fixed_cols, num_fixed = detect_fixed_columns(df)

        # Get sample columns (skip fixed columns and metadata-only columns)
        all_trailing_cols = list(df.columns[num_fixed:])
        sample_cols = [col for col in all_trailing_cols if not self._is_non_sample_column(str(col))]

        # Get Sample_Type row (first data row)
        sample_type_row = df.iloc[0]

        # Build mapping from injection info with better matching
        # Filter out blanks, standards, and non-sample entries
        filtered_injection_list = []
        for info in injection_info_list:
            if self._is_likely_sample_name(info.file_name):
                filtered_injection_list.append(info)

        # Sort by original injection order and re-number starting from 1
        filtered_injection_list.sort(key=lambda x: x.injection_order)
        for new_order, info in enumerate(filtered_injection_list, start=1):
            info.injection_order = new_order

        # Build fast lookup map from normalized keys to injection info
        info_by_key: Dict[str, InjectionInfo] = {}
        for info in filtered_injection_list:
            keys = {
                self._normalize_sample_key(info.file_name),
                self._normalize_sample_key(info.sample_name),
                self._normalize_sample_key(self._simplify_word_sample_name(info.file_name)),
                self._normalize_sample_key(self._extract_primary_sample_token(info.file_name) or ""),
            }
            for key in keys:
                if key and key not in info_by_key:
                    info_by_key[key] = info

        # Build matching data
        sample_info_data = []
        col_to_info_map = {}  # Map column name to matched injection info

        def detect_variant(text: str) -> str:
            """Detect DNA/RNA/DNA+RNA variant from text."""
            text = text.lower().replace("\n", " ").replace("*", " ")
            # Check for DNA+RNA first (must be before individual checks)
            if "dna" in text and "rna" in text and "+" in text:
                return "dna+rna"
            if "dnaandrna" in text.replace(" ", ""):
                return "dna+rna"
            if "_rna" in text or " rna" in text or text.endswith("rna"):
                return "rna"
            return "dna"

        for col in sample_cols:
            col_lower = col.lower()
            col_keys = {
                self._normalize_sample_key(col),
                self._normalize_sample_key(self._extract_primary_sample_token(col) or ""),
            }
            col_keys = {k for k in col_keys if k}

            matched_info: Optional[InjectionInfo] = None
            for key in col_keys:
                if key in info_by_key:
                    matched_info = info_by_key[key]
                    break

            if matched_info is None:
                for info in filtered_injection_list:
                    file_lower = info.file_name.lower().replace("\n", " ")

                    # Match by BC ID with type prefix and variant (DNA/RNA/DNAandRNA)
                    bc_match_col = re.search(r"(tumor|normal|benign|benignfat)?(bc\d+)", col_lower)
                    bc_match_file = re.search(r"(tumor|normal|benign)\s*(tissue)?\s*(fat\s*)?(bc\d+)", file_lower)

                    if bc_match_col and bc_match_file:
                        col_prefix = bc_match_col.group(1) or ""
                        col_id = bc_match_col.group(2)
                        col_variant = detect_variant(col_lower)

                        file_prefix = bc_match_file.group(1) or ""
                        file_id = bc_match_file.group(4)
                        file_variant = detect_variant(file_lower)

                        # Normalize prefix (benignfat -> benign)
                        if "benign" in col_prefix:
                            col_prefix = "benign"

                        if col_id == file_id and col_prefix == file_prefix and col_variant == file_variant:
                            matched_info = info
                            break

                    # Match by QC number
                    qc_match_col = re.search(r"(pooled_?)?qc_?(\d+)", col_lower)
                    qc_match_file = re.search(r"(pooled_?)?qc_?(\d+)", file_lower)
                    if qc_match_col and qc_match_file:
                        if qc_match_col.group(2) == qc_match_file.group(2):
                            matched_info = info
                            break

            col_to_info_map[col] = matched_info

            # Get sample type from the Sample_Type row
            sample_type = sample_type_row.get(col, "sample")

            sample_info_data.append({
                "Sample_Name": matched_info.file_name if matched_info else col,
                "Sample_Type": sample_type,
                "Injection_Order": matched_info.injection_order if matched_info else 999,
                "Injection_Volume": matched_info.injection_volume if matched_info else 0,
                "_col_name": col,  # Internal: for column reordering
            })

        # Create DataFrame and sort by Injection_Order
        sample_info_df = pd.DataFrame(sample_info_data)
        sample_info_df = sample_info_df.sort_values("Injection_Order").reset_index(drop=True)

        # Ensure SampleInfo headers exist and follow expected order
        display_cols = [
            "Sample_Name",
            "Sample_Type",
            "Injection_Order",
            "Batch",
            "Injection_Volume",
            "DNA_mg/20uL",
        ]
        for col in display_cols:
            if col not in sample_info_df.columns:
                sample_info_df[col] = np.nan

        extra_cols = [c for c in sample_info_df.columns if c not in display_cols and c != "_col_name"]
        ordered_cols = display_cols + extra_cols
        if "_col_name" in sample_info_df.columns:
            ordered_cols.append("_col_name")
        sample_info_df = sample_info_df[ordered_cols]

        return sample_info_df

    def _reorder_columns_by_injection(
        self,
        df: pd.DataFrame,
        sample_info_df: pd.DataFrame,
    ) -> pd.DataFrame:
        """
        Reorder DataFrame columns based on Injection_Order from SampleInfo.

        Args:
            df: DataFrame with sample columns
            sample_info_df: SampleInfo DataFrame with Injection_Order

        Returns:
            DataFrame with columns reordered by Injection_Order
        """
        if sample_info_df.empty:
            return df

        # Determine fixed columns
        fixed_cols, num_fixed = detect_fixed_columns(df)

        # Split trailing columns into sample columns and metadata-only columns.
        trailing_cols = list(df.columns[num_fixed:])
        sample_cols = [col for col in trailing_cols if not self._is_non_sample_column(str(col))]
        metadata_cols = [col for col in trailing_cols if self._is_non_sample_column(str(col))]

        # Use _col_name from SampleInfo if available (direct mapping)
        if "_col_name" in sample_info_df.columns:
            # Sort by Injection_Order and get column names in order
            sorted_info = sample_info_df.sort_values("Injection_Order")
            ordered_sample_cols = sorted_info["_col_name"].tolist()

            # Filter to only include columns that exist in df
            ordered_sample_cols = [c for c in ordered_sample_cols if c in sample_cols]

            # Add any remaining columns that weren't in SampleInfo
            for col in sample_cols:
                if col not in ordered_sample_cols:
                    ordered_sample_cols.append(col)
        else:
            # Fallback to matching by sample name
            ordered_sample_cols = []
            sorted_info = sample_info_df.sort_values("Injection_Order")

            for _, row in sorted_info.iterrows():
                sample_name = row["Sample_Name"]

                # Find matching column
                for col in sample_cols:
                    col_lower = col.lower()
                    sample_lower = sample_name.lower()

                    matched = False

                    # BC ID match with type
                    bc_match_col = re.search(r"(tumor|normal|benign|benignfat)?(bc\d+)", col_lower)
                    bc_match_sample = re.search(r"(tumor|normal|benign)?\s*tissue\s*(fat\s*)?(bc\d+)", sample_lower)
                    if bc_match_col and bc_match_sample:
                        col_type = bc_match_col.group(1) or ""
                        if "benign" in col_type:
                            col_type = "benign"
                        sample_type = bc_match_sample.group(1) or ""
                        col_id = bc_match_col.group(2)
                        sample_id = bc_match_sample.group(3)
                        if col_id == sample_id and col_type == sample_type:
                            matched = True

                    # QC match
                    qc_match_col = re.search(r"(pooled_?)?qc_?(\d+)", col_lower)
                    qc_match_sample = re.search(r"(pooled_?)?qc_?(\d+)", sample_lower)
                    if qc_match_col and qc_match_sample:
                        if qc_match_col.group(2) == qc_match_sample.group(2):
                            matched = True

                    if matched and col not in ordered_sample_cols:
                        ordered_sample_cols.append(col)
                        break

            # Add any remaining columns that weren't matched
            for col in sample_cols:
                if col not in ordered_sample_cols:
                    ordered_sample_cols.append(col)

        # Rebuild DataFrame with new column order
        new_column_order = fixed_cols + ordered_sample_cols + metadata_cols
        return df[new_column_order]

    def _parse_method_file(self, file_path: Union[str, Path]) -> Dict[str, str]:
        """
        Parse method file (Word document) to extract sample type mapping.

        Args:
            file_path: Path to the Word document

        Returns:
            Dictionary mapping sample names to sample types
        """
        mapping = {}
        file_path = Path(file_path)

        if not file_path.exists():
            return mapping

        if file_path.suffix.lower() not in [".docx", ".doc"]:
            return mapping

        tables: List[List[List[str]]] = []
        try:
            from docx import Document
            doc = Document(file_path)
            for table in doc.tables:
                table_rows: List[List[str]] = []
                for row in table.rows:
                    table_rows.append([cell.text.strip() for cell in row.cells])
                if table_rows:
                    tables.append(table_rows)
        except ImportError:
            logger.warning("python-docx not installed; using fallback DOCX parser for method file")
            tables = self._extract_docx_tables_fallback(file_path)
        except Exception as exc:
            logger.warning("python-docx parse failed for %s (%s); using fallback parser", file_path, exc)
            tables = self._extract_docx_tables_fallback(file_path)

        # Parse tables for sample information
        for table_rows in tables:
            for row in table_rows:
                cells = [str(cell).strip() for cell in row]

                # Look for patterns like "Tumor tissue BC2257_DNA" or "Normal tissue BC2257_DNA"
                for cell_text in cells:
                    cell_lower = cell_text.lower()

                    # Detect sample type from cell content
                    if "tumor" in cell_lower or "cancer" in cell_lower:
                        sample_id = self._extract_sample_id(cell_text)
                        if sample_id:
                            mapping[sample_id] = "tumor"

                    elif "normal" in cell_lower:
                        sample_id = self._extract_sample_id(cell_text)
                        if sample_id:
                            mapping[sample_id] = "normal"

                    elif "benign" in cell_lower:
                        sample_id = self._extract_sample_id(cell_text)
                        if sample_id:
                            mapping[sample_id] = "benign"

                    elif "qc" in cell_lower or "pool" in cell_lower:
                        sample_id = self._extract_sample_id(cell_text)
                        if sample_id:
                            mapping[sample_id] = "qc"

                    elif "blank" in cell_lower:
                        sample_id = self._extract_sample_id(cell_text)
                        if sample_id:
                            mapping[sample_id] = "blank"

                    elif "std" in cell_lower:
                        sample_id = self._extract_sample_id(cell_text)
                        if sample_id:
                            mapping[sample_id] = "standard"

        return mapping

    def _extract_sample_id(self, text: str) -> Optional[str]:
        """
        Extract sample ID from text.

        Examples:
            "Tumor tissue BC2257_DNA" -> "BC2257_DNA"
            "Normal tissue BC2257_DNA" -> "BC2257_DNA"
        """
        # Pattern to match sample IDs like BC2257_DNA
        patterns = [
            r"BC\d+_\w+",  # BC2257_DNA, BC2257_RNA
            r"[A-Z]{2,}\d+",  # Generic ID pattern
        ]

        for pattern in patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                return match.group()

        return None

    def auto_detect_sample_types(
        self,
        column_names: List[str],
        patterns: Optional[Dict[str, str]] = None,
    ) -> Dict[str, str]:
        """
        Auto-detect sample types from column names.

        Args:
            column_names: List of column names to analyze
            patterns: Custom patterns for detection

        Returns:
            Dictionary mapping column names to detected sample types
        """
        mapping = {}

        for col in column_names:
            sample_type = self._detect_sample_type(col, patterns or {})
            mapping[col] = sample_type

        return mapping


def load_raw_data(
    file_path: Union[str, Path],
    **kwargs,
) -> pd.DataFrame:
    """
    Load raw mass spectrometry data from file.

    Supports TSV, CSV, and Excel formats.

    Args:
        file_path: Path to the data file
        **kwargs: Additional arguments passed to pandas read functions

    Returns:
        DataFrame with raw data
    """
    file_path = Path(file_path)

    if file_path.suffix.lower() == ".tsv":
        return pd.read_csv(file_path, sep="\t", **kwargs)
    elif file_path.suffix.lower() == ".csv":
        return pd.read_csv(file_path, **kwargs)
    elif file_path.suffix.lower() in [".xlsx", ".xls"]:
        return pd.read_excel(file_path, **kwargs)
    else:
        # Try TSV as default
        return pd.read_csv(file_path, sep="\t", **kwargs)
